import io
from typing import List, Dict, Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def extract_docx_text(file_like: io.BytesIO) -> str:
    file_like.seek(0)
    doc = Document(file_like)
    texts: List[str] = []
    for p in doc.paragraphs:
        texts.append(p.text)
    return "\n".join(texts)


def annotate_docx_with_issues(file_like: io.BytesIO, issues: List[Dict[str, Any]]) -> bytes:
    """
    Apply simple inline highlights and append a 'Review Comments' section.
    python-docx has no official comments API; this approach uses highlighting and
    appended notes to simulate reviewer comments.
    """
    file_like.seek(0)
    doc = Document(file_like)

    # naive highlight: for each issue, highlight first paragraph containing keyword pattern
    for issue in issues:
        pattern = issue.get("pattern", None)
        if not pattern:
            continue
        for para in doc.paragraphs:
            if pattern.lower() in para.text.lower():
                # highlight all runs in this paragraph
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                # add inline bracketed comment
                inline_parts = [issue['issue']]
                if issue.get('suggestion'):
                    inline_parts.append(f"Suggestion: {issue['suggestion']}")
                if issue.get('citation'):
                    inline_parts.append(f"Citation: {issue['citation']}")
                para.add_run(" [COMMENT: " + " | ".join(inline_parts) + "] ")
                break

    # Append a review section (avoid style names that may be missing)
    try:
        doc.add_page_break()
    except Exception:
        pass
    try:
        doc.add_heading("Review Comments (Corporate Agent)", level=1)
    except Exception:
        title_p = doc.add_paragraph("Review Comments (Corporate Agent)")
        if title_p.runs:
            title_p.runs[0].bold = True
    for idx, issue in enumerate(issues, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. Section: {issue.get('section','N/A')}\n").bold = True
        doc.add_paragraph(f"Issue: {issue['issue']}")
        doc.add_paragraph(f"Severity: {issue.get('severity','N/A')}")
        if issue.get("suggestion"):
            doc.add_paragraph(f"Suggestion: {issue['suggestion']}")
        if issue.get("citation"):
            doc.add_paragraph(f"Citation: {issue['citation']}")

    out_bytes = io.BytesIO()
    doc.save(out_bytes)
    return out_bytes.getvalue()


