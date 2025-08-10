"""
Utility to generate a sample .docx (before) and run the analyzer to produce a reviewed file (after).

Usage:
  python scripts/make_sample_docs.py

Outputs:
  - sample_docs/example_before.docx
  - outputs/example_before__reviewed.docx
  - outputs/analysis_summary.json (overwritten)
"""

import io
import os
from typing import List, Dict, Any

from docx import Document

from document_checker import analyze_documents


def ensure_dirs() -> None:
    os.makedirs("sample_docs", exist_ok=True)
    os.makedirs("outputs", exist_ok=True)


def write_sample_docx(path: str) -> None:
    doc = Document()
    doc.add_heading("Articles of Association", level=1)
    doc.add_paragraph("This document sets out the rules of the Company.")
    doc.add_paragraph("Governing law and jurisdiction: UAE Federal Courts.")
    doc.add_paragraph("The Company may, from time to time, take actions as appropriate.")
    doc.add_paragraph("Shareholding and capital provisions are to be determined.")
    # No signature block on purpose to trigger an issue
    doc.save(path)


def main() -> None:
    ensure_dirs()
    sample_path = os.path.join("sample_docs", "example_before.docx")
    write_sample_docx(sample_path)

    with open(sample_path, "rb") as f:
        uploaded_docs: List[Dict[str, Any]] = [{
            "filename": os.path.basename(sample_path),
            "bytes": f.read(),
        }]

    # Run analysis without forcing LLM JSON to avoid API requirements
    _ = analyze_documents(uploaded_docs, rag_index=None, run_llm_json=False)
    print("Sample generated:")
    print(" - before:", sample_path)
    print(" - after :", os.path.join("outputs", "example_before__reviewed.docx"))
    print(" - summary:", os.path.join("outputs", "analysis_summary.json"))


if __name__ == "__main__":
    main()


